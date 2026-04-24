"""Guidelines RAG - reads design-standard documents and injects relevant chunks into prompts.

Supported document formats: HTML, PDF, DOCX, TXT, MD.

Pipeline position: Stage 0.5 - runs AFTER structural pre-check, BEFORE elaboration.

Embedding strategy:
  Primary  : Ollama /api/embeddings (e.g. nomic-embed-text) - semantic vector search
  Fallback : TF-IDF keyword overlap (pure Python) when Ollama is unreachable

Cache: per-document JSON files in guidelines_cache_dir, keyed by SHA-256 of file content.
       Cache is invalidated automatically when the source file changes.
"""

from __future__ import annotations

import hashlib
import json
import logging
import math
import re
import time
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

import httpx

logger = logging.getLogger(__name__)

# Ollama embedding endpoint
_OLLAMA_BASE = "http://localhost:11434"
_EMBED_ENDPOINT = "/api/embeddings"

# Per-diagram-type vocabulary boosts retrieval relevance
_DIAGRAM_VOCAB: dict[str, str] = {
    "activity":      "activity flowchart decision fork join action runnable cycle",
    "sequence":      "sequence message flow actor lifeline interaction call",
    "state_machine": "state machine transition guard condition trigger event mode",
    "class":         "class diagram attribute operation method inheritance interface",
    "component":     "component interface port connection provided required SWC",
}

# Supported file extensions
_SUPPORTED_EXTS = {".html", ".htm", ".pdf", ".docx", ".txt", ".md", ".xlsx", ".xls"}


# ── Data Classes ──────────────────────────────────────────────────────────────

@dataclass
class DocumentChunk:
    """One chunk of text extracted from a guideline document."""
    id: int
    section: str                              # heading/title or filename stem
    text: str                                 # up to chunk_size chars
    source_file: str                          # originating filename
    embedding: Optional[list[float]] = None  # None = keyword fallback in use


@dataclass
class GuidelinesStatus:
    """Summary returned after load_all() and exposed via /guidelines/status."""
    enabled: bool
    doc_count: int
    chunk_count: int
    embedded: bool         # True = Ollama vector embeddings; False = keyword TF-IDF
    files: list[str] = field(default_factory=list)
    cache_hits: int = 0
    message: str = ""

    def to_dict(self) -> dict:
        return {
            "enabled": self.enabled,
            "doc_count": self.doc_count,
            "chunk_count": self.chunk_count,
            "embedded": self.embedded,
            "files": self.files,
            "cache_hits": self.cache_hits,
            "message": self.message,
        }


# ── Main class ────────────────────────────────────────────────────────────────

class GuidelinesReader:
    """Reads, chunks, embeds, caches, and retrieves design guideline documents.

    Usage:
        reader = GuidelinesReader(settings)
        status = await reader.load_all(progress_callback=cb)
        block  = await reader.build_guidelines_context("activity", req_text)
    """

    def __init__(self, settings) -> None:
        self._settings = settings
        self._guidelines_dir: Path = settings.get_guidelines_dir()
        self._cache_dir: Path = settings.get_guidelines_cache_dir()
        self._chunks_by_file: dict[str, list[DocumentChunk]] = {}
        self._all_chunks: list[DocumentChunk] = []
        self._embedded: bool = False
        self._status: Optional[GuidelinesStatus] = None

    # ── Public API ────────────────────────────────────────────────────────────

    async def load_all(
        self,
        progress_callback: Optional[callable] = None,
    ) -> GuidelinesStatus:
        """Scan guidelines_dir, parse docs, embed chunks (or fallback), cache results.

        Returns a GuidelinesStatus with counts. Also emits a 'guidelines' SSE event
        via progress_callback if provided.
        """
        if not self._settings.guidelines_enabled:
            status = GuidelinesStatus(enabled=False, doc_count=0, chunk_count=0, embedded=False,
                                      message="[Guidelines] Disabled (MUD_GUIDELINES_ENABLED=false)")
            self._status = status
            return status

        if not self._guidelines_dir.exists():
            logger.info("[Guidelines] Directory not found: %s", self._guidelines_dir)
            status = GuidelinesStatus(enabled=True, doc_count=0, chunk_count=0, embedded=False,
                                      message=f"[Guidelines] No guidelines dir at {self._guidelines_dir}")
            self._status = status
            return status

        files = self._get_supported_files()
        if not files:
            status = GuidelinesStatus(enabled=True, doc_count=0, chunk_count=0, embedded=False,
                                      files=[], message="[Guidelines] No supported documents found")
            self._status = status
            return status

        self._cache_dir.mkdir(parents=True, exist_ok=True)
        cache_hits = 0
        all_chunks: list[DocumentChunk] = []

        for path in files:
            file_hash = self._file_hash(path)
            cached = self._load_cache(path, file_hash)
            if cached is not None:
                self._chunks_by_file[path.name] = cached
                all_chunks.extend(cached)
                cache_hits += 1
                logger.info("[Guidelines] Cache hit: %s (%d chunks)", path.name, len(cached))
                continue

            # Parse + chunk
            try:
                sections = self._extract_text(path)
            except Exception as exc:
                logger.warning("[Guidelines] Failed to parse %s: %s", path.name, exc)
                continue

            chunk_id = len(all_chunks)
            chunks: list[DocumentChunk] = []
            for section_title, section_text in sections:
                new_chunks = self._chunk_section(
                    section=section_title,
                    text=section_text,
                    source_file=path.name,
                    chunk_size=self._settings.guidelines_chunk_size,
                    chunk_id_start=chunk_id,
                )
                chunks.extend(new_chunks)
                chunk_id += len(new_chunks)

            if not chunks:
                continue

            # Attempt Ollama embeddings
            chunks, used_vectors = await self._embed_chunks(chunks)

            self._save_cache(path, file_hash, chunks)
            self._chunks_by_file[path.name] = chunks
            all_chunks.extend(chunks)
            logger.info("[Guidelines] Parsed %s: %d chunks (embedded=%s)",
                        path.name, len(chunks), used_vectors)

        self._all_chunks = all_chunks
        self._embedded = any(c.embedding is not None for c in all_chunks)

        embed_label = "vector" if self._embedded else "keyword"
        msg = (
            f"[Guidelines] {len(files)} doc(s), {len(all_chunks)} chunks "
            f"({embed_label} retrieval)"
        )
        status = GuidelinesStatus(
            enabled=True,
            doc_count=len(files),
            chunk_count=len(all_chunks),
            embedded=self._embedded,
            files=[p.name for p in files],
            cache_hits=cache_hits,
            message=msg,
        )
        self._status = status

        if progress_callback:
            progress_callback({
                "stage": "guidelines",
                "doc_count": len(files),
                "chunk_count": len(all_chunks),
                "embedded": self._embedded,
                "files": [p.name for p in files],
                "message": msg,
            })

        return status

    async def build_guidelines_context(
        self,
        diagram_type: str,
        requirements_text: str,
    ) -> str:
        """Retrieve top-N most relevant chunks and return a Markdown context block.

        Returns empty string if no guidelines are loaded.
        """
        if not self._all_chunks:
            return ""

        vocab = _DIAGRAM_VOCAB.get(diagram_type, diagram_type)
        query = f"{vocab} {requirements_text[:500]}"

        top_chunks = await self._retrieve_top_n(
            query_text=query,
            all_chunks=self._all_chunks,
            n=self._settings.guidelines_max_chunks,
        )

        if not top_chunks:
            return ""

        return self._format_context_block(top_chunks)

    def clear_cache(self) -> int:
        """Delete all cached chunk JSON files. Returns count of files deleted."""
        if not self._cache_dir.exists():
            return 0
        deleted = 0
        for f in self._cache_dir.glob("*.json"):
            try:
                f.unlink()
                deleted += 1
            except OSError:
                pass
        logger.info("[Guidelines] Cleared %d cache file(s)", deleted)
        return deleted

    def get_status(self) -> dict:
        """Fast status scan (no embedding load). For the /guidelines/status endpoint."""
        if not self._settings.guidelines_enabled:
            return {
                "enabled": False,
                "doc_count": 0, "chunk_count": 0, "embedded": False,
                "files": [], "guidelines_dir": str(self._guidelines_dir),
                "cache_dir": str(self._cache_dir),
            }

        files = self._get_supported_files() if self._guidelines_dir.exists() else []
        # Count cached chunks without loading embeddings
        chunk_count = 0
        has_embeddings = False
        for f in files:
            cache_path = self._cache_file_path(f)
            if cache_path.exists():
                try:
                    data = json.loads(cache_path.read_text(encoding="utf-8"))
                    chunk_count += len(data.get("chunks", []))
                    if any(c.get("embedding") for c in data.get("chunks", [])):
                        has_embeddings = True
                except Exception:
                    pass

        return {
            "enabled": True,
            "doc_count": len(files),
            "chunk_count": chunk_count,
            "embedded": has_embeddings,
            "files": [f.name for f in files],
            "guidelines_dir": str(self._guidelines_dir),
            "cache_dir": str(self._cache_dir),
        }

    # ── File helpers ──────────────────────────────────────────────────────────

    def _get_supported_files(self) -> list[Path]:
        if not self._guidelines_dir.exists():
            return []
        return sorted(
            p for p in self._guidelines_dir.iterdir()
            if p.is_file() and p.suffix.lower() in _SUPPORTED_EXTS
            and not p.name.startswith(".")
        )

    def _file_hash(self, path: Path) -> str:
        """SHA-256 of first 64KB of file (fast, sufficient for cache keying)."""
        h = hashlib.sha256()
        with path.open("rb") as f:
            h.update(f.read(65536))
        return h.hexdigest()

    def _cache_file_path(self, path: Path) -> Path:
        key = hashlib.sha256(path.name.encode()).hexdigest()[:16]
        return self._cache_dir / f"{key}.json"

    def _load_cache(self, path: Path, file_hash: str) -> Optional[list[DocumentChunk]]:
        cache_path = self._cache_file_path(path)
        if not cache_path.exists():
            return None
        try:
            data = json.loads(cache_path.read_text(encoding="utf-8"))
            if data.get("file_hash") != file_hash:
                return None  # file changed - invalidate
            return [
                DocumentChunk(
                    id=c["id"],
                    section=c["section"],
                    text=c["text"],
                    source_file=c.get("source_file", path.name),
                    embedding=c.get("embedding"),
                )
                for c in data.get("chunks", [])
            ]
        except Exception as exc:
            logger.debug("[Guidelines] Cache read error for %s: %s", path.name, exc)
            return None

    def _save_cache(self, path: Path, file_hash: str, chunks: list[DocumentChunk]) -> None:
        cache_path = self._cache_file_path(path)
        data = {
            "filename": path.name,
            "file_hash": file_hash,
            "embed_model": self._settings.guidelines_embed_model,
            "cached_at": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
            "chunks": [
                {
                    "id": c.id,
                    "section": c.section,
                    "text": c.text,
                    "source_file": c.source_file,
                    "embedding": c.embedding,
                }
                for c in chunks
            ],
        }
        try:
            cache_path.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
        except Exception as exc:
            logger.warning("[Guidelines] Failed to save cache for %s: %s", path.name, exc)

    # ── Document parsers ──────────────────────────────────────────────────────

    def _extract_text(self, path: Path) -> list[tuple[str, str]]:
        """Parse document → list of (section_title, text) pairs."""
        suffix = path.suffix.lower()
        if suffix in {".html", ".htm"}:
            return self._parse_html(path)
        elif suffix == ".pdf":
            return self._parse_pdf(path)
        elif suffix == ".docx":
            return self._parse_docx(path)
        elif suffix in {".xlsx", ".xls"}:
            return self._parse_excel(path)
        else:  # .txt, .md
            return self._parse_plaintext(path)

    def _parse_html(self, path: Path) -> list[tuple[str, str]]:
        """Extract sections from HTML by walking h1/h2/h3 heading tags."""
        try:
            from lxml import html as lxml_html
            raw = path.read_bytes()
            doc = lxml_html.fromstring(raw)

            # Strip script and style elements
            for tag in doc.cssselect("script, style, noscript"):
                tag.getparent().remove(tag)

            sections: list[tuple[str, str]] = []
            current_title = path.stem
            current_lines: list[str] = []

            for el in doc.iter():
                tag = el.tag if isinstance(el.tag, str) else ""
                if tag in {"h1", "h2", "h3"}:
                    if current_lines:
                        text = " ".join(current_lines).strip()
                        if len(text) > 40:
                            sections.append((current_title, text))
                    current_title = (el.text_content() or "").strip()
                    current_lines = []
                else:
                    text = (el.text or "").strip()
                    if text and len(text) > 10:
                        current_lines.append(text)

            if current_lines:
                text = " ".join(current_lines).strip()
                if len(text) > 40:
                    sections.append((current_title, text))

            if not sections:
                # Fallback: full text content
                full = doc.text_content().strip()
                if full:
                    sections = [(path.stem, full)]

            return sections
        except Exception as exc:
            logger.warning("[Guidelines] HTML parse error %s: %s", path.name, exc)
            return []

    def _parse_pdf(self, path: Path) -> list[tuple[str, str]]:
        """Extract sections from PDF using pypdf."""
        try:
            import pypdf
            reader = pypdf.PdfReader(str(path))
            sections: list[tuple[str, str]] = []
            current_title = path.stem
            current_lines: list[str] = []

            _ALL_CAPS = re.compile(r'^[A-Z][A-Z0-9 /&\-]{4,}$')

            for page in reader.pages:
                page_text = page.extract_text() or ""
                for line in page_text.splitlines():
                    line = line.strip()
                    if not line:
                        continue
                    if _ALL_CAPS.match(line) and len(line) < 80:
                        if current_lines:
                            text = " ".join(current_lines).strip()
                            if len(text) > 40:
                                sections.append((current_title, text))
                        current_title = line.title()
                        current_lines = []
                    else:
                        current_lines.append(line)

            if current_lines:
                text = " ".join(current_lines).strip()
                if len(text) > 40:
                    sections.append((current_title, text))

            if not sections:
                full = " ".join(
                    (p.extract_text() or "") for p in reader.pages
                ).strip()
                sections = [(path.stem, full)]

            return sections
        except ImportError:
            logger.warning("[Guidelines] pypdf not installed - cannot parse PDF files")
            return []
        except Exception as exc:
            logger.warning("[Guidelines] PDF parse error %s: %s", path.name, exc)
            return []

    def _parse_docx(self, path: Path) -> list[tuple[str, str]]:
        """Extract sections from DOCX using python-docx."""
        try:
            import docx  # python-docx
            doc = docx.Document(str(path))
            sections: list[tuple[str, str]] = []
            current_title = path.stem
            current_lines: list[str] = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                style_name = para.style.name if para.style else ""
                if any(h in style_name for h in ("Heading 1", "Heading 2", "Heading 3")):
                    if current_lines:
                        body = " ".join(current_lines).strip()
                        if len(body) > 40:
                            sections.append((current_title, body))
                    current_title = text
                    current_lines = []
                else:
                    current_lines.append(text)

            if current_lines:
                body = " ".join(current_lines).strip()
                if len(body) > 40:
                    sections.append((current_title, body))

            if not sections:
                full = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                sections = [(path.stem, full)]

            return sections
        except ImportError:
            logger.warning("[Guidelines] python-docx not installed - cannot parse DOCX files")
            return []
        except Exception as exc:
            logger.warning("[Guidelines] DOCX parse error %s: %s", path.name, exc)
            return []

    def _parse_excel(self, path: Path) -> list[tuple[str, str]]:
        """Extract sections from Excel by treating each sheet as a section.

        Each row is converted to 'col1: val1 | col2: val2 ...' text so the
        content is searchable by the RAG retriever.
        """
        try:
            import openpyxl
            wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
            sections: list[tuple[str, str]] = []

            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                rows = list(ws.iter_rows(values_only=True))
                if not rows:
                    continue

                # First row as headers (fall back to column letters if empty)
                headers = [str(h).strip() if h else f"Col{i}" for i, h in enumerate(rows[0])]

                lines: list[str] = []
                for row in rows[1:]:
                    parts = []
                    for header, cell in zip(headers, row):
                        if cell is not None and str(cell).strip():
                            parts.append(f"{header}: {str(cell).strip()}")
                    if parts:
                        lines.append(" | ".join(parts))

                if lines:
                    body = "\n".join(lines)
                    sections.append((f"{path.stem} — {sheet_name}", body))

            wb.close()
            if not sections:
                sections = [(path.stem, "Empty workbook")]
            return sections
        except ImportError:
            logger.warning("[Guidelines] openpyxl not installed - cannot parse Excel files")
            return []
        except Exception as exc:
            logger.warning("[Guidelines] Excel parse error %s: %s", path.name, exc)
            return []

    def _parse_plaintext(self, path: Path) -> list[tuple[str, str]]:
        """Extract sections from TXT/MD by Markdown headings or paragraph breaks."""
        try:
            text = path.read_text(encoding="utf-8", errors="replace")
            sections: list[tuple[str, str]] = []
            current_title = path.stem
            current_lines: list[str] = []

            _MD_HEADING = re.compile(r'^#{1,3}\s+(.+)$')

            for line in text.splitlines():
                m = _MD_HEADING.match(line)
                if m:
                    if current_lines:
                        body = " ".join(current_lines).strip()
                        if len(body) > 40:
                            sections.append((current_title, body))
                    current_title = m.group(1).strip()
                    current_lines = []
                else:
                    stripped = line.strip()
                    if stripped:
                        current_lines.append(stripped)
                    elif current_lines and len(" ".join(current_lines)) > 200:
                        # Long paragraph - emit as own chunk
                        body = " ".join(current_lines).strip()
                        sections.append((current_title, body))
                        current_lines = []

            if current_lines:
                body = " ".join(current_lines).strip()
                if len(body) > 40:
                    sections.append((current_title, body))

            if not sections:
                sections = [(path.stem, text.strip())]

            return sections
        except Exception as exc:
            logger.warning("[Guidelines] Text parse error %s: %s", path.name, exc)
            return []

    # ── Chunking ──────────────────────────────────────────────────────────────

    def _chunk_section(
        self,
        section: str,
        text: str,
        source_file: str,
        chunk_size: int,
        chunk_id_start: int,
    ) -> list[DocumentChunk]:
        """Split section text into sentence-boundary chunks of <= chunk_size chars."""
        if not text.strip():
            return []

        # Split into sentences (naive - split on '. ' boundary)
        sentences = re.split(r'(?<=[.!?])\s+', text.strip())
        chunks: list[DocumentChunk] = []
        buffer: list[str] = []
        buffer_len = 0
        chunk_idx = chunk_id_start

        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
            if buffer_len + len(sentence) + 1 > chunk_size and buffer:
                chunks.append(DocumentChunk(
                    id=chunk_idx,
                    section=section,
                    text=" ".join(buffer),
                    source_file=source_file,
                ))
                chunk_idx += 1
                buffer = []
                buffer_len = 0
            # If a single sentence exceeds chunk_size, split it hard
            if len(sentence) > chunk_size:
                for i in range(0, len(sentence), chunk_size):
                    part = sentence[i: i + chunk_size]
                    chunks.append(DocumentChunk(
                        id=chunk_idx,
                        section=section,
                        text=part,
                        source_file=source_file,
                    ))
                    chunk_idx += 1
            else:
                buffer.append(sentence)
                buffer_len += len(sentence) + 1

        if buffer:
            chunks.append(DocumentChunk(
                id=chunk_idx,
                section=section,
                text=" ".join(buffer),
                source_file=source_file,
            ))

        return chunks

    # ── Embedding ─────────────────────────────────────────────────────────────

    async def _embed_chunks(
        self, chunks: list[DocumentChunk]
    ) -> tuple[list[DocumentChunk], bool]:
        """Attempt Ollama embeddings for all chunks. Falls back to None on failure."""
        model = self._settings.guidelines_embed_model
        success_count = 0

        try:
            async with httpx.AsyncClient(timeout=20) as client:
                for chunk in chunks:
                    try:
                        resp = await client.post(
                            f"{_OLLAMA_BASE}{_EMBED_ENDPOINT}",
                            json={"model": model, "prompt": chunk.text},
                        )
                        if resp.status_code == 200:
                            chunk.embedding = resp.json().get("embedding")
                            if chunk.embedding:
                                success_count += 1
                    except (httpx.TimeoutException, httpx.HTTPStatusError):
                        pass  # leave embedding=None
        except httpx.ConnectError:
            logger.info(
                "[Guidelines] Ollama not reachable at %s - using keyword fallback", _OLLAMA_BASE
            )
            return chunks, False

        used_vectors = success_count > 0
        if used_vectors:
            logger.info(
                "[Guidelines] Embedded %d/%d chunks via %s",
                success_count, len(chunks), model,
            )
        else:
            logger.info("[Guidelines] No embeddings obtained - keyword fallback active")
        return chunks, used_vectors

    async def _embed_query(self, query_text: str) -> Optional[list[float]]:
        """Embed a single query string via Ollama. Returns None on failure."""
        model = self._settings.guidelines_embed_model
        try:
            async with httpx.AsyncClient(timeout=10) as client:
                resp = await client.post(
                    f"{_OLLAMA_BASE}{_EMBED_ENDPOINT}",
                    json={"model": model, "prompt": query_text},
                )
                if resp.status_code == 200:
                    return resp.json().get("embedding")
        except Exception:
            pass
        return None

    # ── Retrieval ─────────────────────────────────────────────────────────────

    async def _retrieve_top_n(
        self,
        query_text: str,
        all_chunks: list[DocumentChunk],
        n: int,
    ) -> list[DocumentChunk]:
        """Score all chunks against query and return top-N deduplicated by section."""
        if not all_chunks:
            return []

        scored: list[tuple[float, DocumentChunk]] = []

        if self._embedded:
            query_emb = await self._embed_query(query_text)
            if query_emb:
                for chunk in all_chunks:
                    if chunk.embedding:
                        score = self._cosine_similarity(query_emb, chunk.embedding)
                    else:
                        score = self._tfidf_score(query_text, chunk.text) * 0.5
                    scored.append((score, chunk))
            else:
                # Ollama down during retrieval - fallback to keyword
                for chunk in all_chunks:
                    scored.append((self._tfidf_score(query_text, chunk.text), chunk))
        else:
            for chunk in all_chunks:
                scored.append((self._tfidf_score(query_text, chunk.text), chunk))

        # Sort descending
        scored.sort(key=lambda x: x[0], reverse=True)

        # Deduplicate: keep best per (source_file, section) pair
        seen: set[tuple[str, str]] = set()
        top: list[DocumentChunk] = []
        for score, chunk in scored:
            key = (chunk.source_file, chunk.section)
            if key not in seen:
                seen.add(key)
                top.append(chunk)
            if len(top) >= n:
                break

        return top

    @staticmethod
    def _cosine_similarity(a: list[float], b: list[float]) -> float:
        """Pure-Python cosine similarity."""
        if len(a) != len(b):
            return 0.0
        dot = sum(x * y for x, y in zip(a, b))
        norm_a = math.sqrt(sum(x * x for x in a))
        norm_b = math.sqrt(sum(y * y for y in b))
        if norm_a == 0.0 or norm_b == 0.0:
            return 0.0
        return dot / (norm_a * norm_b)

    @staticmethod
    def _tfidf_score(query_text: str, chunk_text: str) -> float:
        """Token intersection ratio (Jaccard-like) for keyword fallback."""
        def tokenize(t: str) -> set[str]:
            return {w.lower() for w in re.split(r'\W+', t) if len(w) >= 3}

        q_tokens = tokenize(query_text)
        c_tokens = tokenize(chunk_text)
        if not q_tokens or not c_tokens:
            return 0.0
        return len(q_tokens & c_tokens) / len(q_tokens | c_tokens)

    # ── Output formatting ─────────────────────────────────────────────────────

    def _format_context_block(self, chunks: list[DocumentChunk]) -> str:
        """Format retrieved chunks as a Markdown block ready for prompt injection."""
        lines = ["### DESIGN GUIDELINES (from project standards)\n"]
        for chunk in chunks:
            lines.append(f"*Source: {chunk.source_file} - Section: {chunk.section}*\n")
            lines.append(chunk.text[: self._settings.guidelines_chunk_size])
            lines.append("")
        lines.append("---")
        return "\n".join(lines)
